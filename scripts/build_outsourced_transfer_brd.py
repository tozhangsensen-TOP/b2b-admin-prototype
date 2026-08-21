from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "外租库短驳调拨建议业务需求.md"
DOCX_PATH = ROOT / "docs" / "外租库短驳调拨建议业务需求.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 102, 112)
BLACK = RGBColor(0, 0, 0)
LIGHT_FILL = "F2F4F7"
BORDER = "D9E0E7"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_cell_text(cell, bold=False, color=BLACK, size=10.5):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)
    for paragraph in cell.paragraphs:
      set_paragraph_spacing(paragraph, after=0, line=1.10)
      for run in paragraph.runs:
          set_run_font(run, size=size, bold=bold, color=color)


def add_rule(paragraph, color="2E74B5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_title_block(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "业务需求提报"
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "外租库短驳调拨建议业务需求"
    set_run_font(footer.runs[0], size=9, color=MUTED)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    run = p.add_run("外租库短驳调拨建议业务需求")
    set_run_font(run, size=22, bold=True, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=12)
    run = p.add_run("石家庄试点｜面向大数据/IT提报")
    set_run_font(run, size=12, color=MUTED)

    meta = [
        ("版本", "V1.0"),
        ("适用场景", "石家庄本库与外租库短驳调拨"),
        ("需求类型", "业务侧需求说明，不作为页面PRD"),
        ("核心目标", "系统生成调拨建议，人工确认后形成调拨计划/转拨单"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row, (label, value) in zip(table.rows, meta):
        row.cells[0].width = Inches(1.15)
        row.cells[1].width = Inches(5.1)
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], LIGHT_FILL)
        style_cell_text(row.cells[0], bold=True, color=DARK_BLUE)
        style_cell_text(row.cells[1])

    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=8, after=8)
    add_rule(rule)


def add_business_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        hdr[i].width = Inches(widths[i])
        shade_cell(hdr[i], LIGHT_FILL)
        style_cell_text(hdr[i], bold=True, color=DARK_BLUE, size=10)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
            cells[i].width = Inches(widths[i])
            style_cell_text(cells[i], size=10)
    doc.add_paragraph()


def add_para(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        run = p.add_run(part.strip("`") if part.startswith("`") else part)
        set_run_font(run, size=11, color=BLACK)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)
    add_title_block(doc)

    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    in_scope_table = False
    pending_table = []

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("版本：") or line.startswith("适用场景：") or line.startswith("需求类型："):
            continue
        if line.startswith("## "):
            if pending_table:
                add_business_table(doc, pending_table[0], pending_table[1:], [1.2, 5.05])
                pending_table = []
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\. ", line):
            add_number(doc, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif " = " in line and len(line) < 90:
            pending_table.append(["口径", "说明"] if not pending_table else None)
            if pending_table[-1] is None:
                pending_table.pop()
            left, right = line.split(" = ", 1)
            pending_table.append([left, right])
        else:
            if pending_table:
                add_business_table(doc, pending_table[0], pending_table[1:], [1.2, 5.05])
                pending_table = []
            add_para(doc, line)

    if pending_table:
        add_business_table(doc, pending_table[0], pending_table[1:], [1.2, 5.05])

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
